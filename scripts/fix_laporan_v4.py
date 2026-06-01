"""
Fix script: Laporan KP.V3.docx -> Laporan KP.V4.docx
Fixes:
  1. DAFTAR TABLE -> DAFTAR TABEL
  2. TOC double dot 5.. -> 5.
  3. Table -> Tabel (before digit)
  4. 0.x.x table numbering -> x.x
  5. Section heading: 2nd 'Multiclass Logistic Regression' -> 'Deteksi Near Real-Time'
  6. Section heading: 2nd 'Logging dan Security Event' -> 'Alert dan Response Keamanan'
  7. 'Isi gambar:' editorial notes -> academic descriptions
  8. Remove internal editorial note
  9. Remove orphaned sentence
 10. Add Gambar 5.4 description if missing
 11. Add Gambar 5.1-5.4 to DAFTAR GAMBAR list
"""
import re
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"D:\project\Detector\docs\thesis\Laporan KP.V3.docx"
DST = r"D:\project\Detector\docs\thesis\Laporan KP.V4.docx"

# Academic figure descriptions (formal Indonesian)
FIG_DESC = {
    1: (
        "Gambar di atas menampilkan tampilan dashboard monitoring yang menyajikan ringkasan "
        "event keamanan, jumlah alert aktif, tren ancaman dalam periode waktu tertentu, serta "
        "distribusi jenis serangan yang dapat dipantau oleh analis atau operator sistem."
    ),
    2: (
        "Gambar di atas menampilkan halaman detail alert yang memuat jenis ancaman yang "
        "terdeteksi, tingkat severity, timestamp kejadian, actor atau IP sumber, path request, "
        "detector yang menghasilkan alert, dan evidence pendukung yang dapat digunakan analis "
        "untuk menentukan tindak lanjut."
    ),
    3: (
        "Gambar di atas menampilkan hasil evaluasi perbandingan performa metode rules-only, "
        "ML-only, dan hybrid berdasarkan metrik accuracy, precision, recall, macro F1, dan "
        "weighted F1. Hasil evaluasi menunjukkan pendekatan hybrid memiliki performa tertinggi "
        "dibandingkan pendekatan tunggal."
    ),
    4: (
        "Gambar di atas menampilkan halaman pengelolaan model machine learning yang menyajikan "
        "daftar model yang telah dilatih, metadata model aktif, status deployment lock, dan "
        "informasi drift sebagai dasar pemantauan kebutuhan retraining model."
    ),
}

# New DAFTAR GAMBAR entries (inserted in reverse so order comes out correct)
DAFTAR_GAMBAR_NEW = [
    "Gambar 5.4    Monitoring Model dan Drift                                         34",
    "Gambar 5.3    Dashboard Evaluasi Deteksi                                         34",
    "Gambar 5.2    Detail Alert Keamanan                                              32",
    "Gambar 5.1    Dashboard Monitoring Event dan Alert                               31",
]


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def get_text(para):
    return "".join(r.text for r in para.runs)


def set_text(para, new_text):
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for r in para.runs[1:]:
        r.text = ""


def remove_para(para):
    para._p.getparent().remove(para._p)


def insert_para_after(ref_para, text, copy_style_from=None):
    """Insert a plain paragraph with text immediately after ref_para."""
    new_p = OxmlElement("w:p")
    if copy_style_from is not None:
        src_pPr = copy_style_from._p.find(qn("w:pPr"))
        if src_pPr is not None:
            new_p.append(copy.deepcopy(src_pPr))
    new_r = OxmlElement("w:r")
    new_t = OxmlElement("w:t")
    new_t.text = text
    new_r.append(new_t)
    new_p.append(new_r)
    ref_para._p.addnext(new_p)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
doc = Document(SRC)
paras = list(doc.paragraphs)

to_remove = set()
multiclass_lr_count = 0
logging_security_count = 0
current_fig = None
gambar54_para = None
gambar54_has_desc = False
gambar42_list_para = None

changes = []


# ===========================================================================
# PASS 1 — text replacements & identification
# ===========================================================================
for i, para in enumerate(paras):
    text = get_text(para)
    if not text.strip():
        continue

    new_text = text

    # 1. DAFTAR TABLE -> DAFTAR TABEL
    new_text = new_text.replace("DAFTAR TABLE", "DAFTAR TABEL")

    # 2. TOC double dot  5.. -> 5.
    new_text = re.sub(r"5\.\.(\d)", r"5.\1", new_text)

    # 3. Table -> Tabel  (before digit)
    new_text = re.sub(r"\bTable (\d)", r"Tabel \1", new_text)

    # 4. 0.x.x table numbering (applied after step 3 so 'Tabel' already set)
    number_map = [
        (r"\bTabel 0\.1\.1\b", "Tabel 1.1"),
        (r"\bTabel 3\.0\.1\b", "Tabel 3.1"),
        (r"\bTabel 3\.0\.2\b", "Tabel 3.2"),
        (r"\bTabel 4\.0\.1\b", "Tabel 4.1"),
        (r"\bTabel 4\.0\.2\b", "Tabel 4.2"),
        (r"\bTabel 4\.0\.7\b", "Tabel 4.7"),
        (r"\bTabel 5\.0\.2\b", "Tabel 5.2"),
        (r"\bTabel 5\.0\.3\b", "Tabel 5.3"),
        (r"\bTabel 5\.0\.4\b", "Tabel 5.4"),
        (r"\bTabel 5\.0\.5\b", "Tabel 5.5"),
    ]
    for pat, rep in number_map:
        new_text = re.sub(pat, rep, new_text)

    if new_text != text:
        set_text(para, new_text)
        text = new_text
        changes.append(f"  TEXT [{i+1:3d}]: {text[:80]!r}")

    # 5. Second 'Multiclass Logistic Regression' heading -> 'Deteksi Near Real-Time'
    if text.strip() == "Multiclass Logistic Regression":
        multiclass_lr_count += 1
        if multiclass_lr_count == 2:
            set_text(para, "Deteksi Near Real-Time")
            changes.append(f"  HEAD [{i+1:3d}]: 'Multiclass LR' #2 -> 'Deteksi Near Real-Time'")

    # 6. Second 'Logging dan Security Event' heading -> 'Alert dan Response Keamanan'
    if text.strip() == "Logging dan Security Event":
        logging_security_count += 1
        if logging_security_count == 2:
            set_text(para, "Alert dan Response Keamanan")
            changes.append(f"  HEAD [{i+1:3d}]: 'Logging...' #2 -> 'Alert dan Response Keamanan'")

    # 7. Track figure number for 5.x
    m = re.match(r"Gambar 5\.(\d+)", text)
    if m:
        current_fig = int(m.group(1))
        if current_fig == 4:
            gambar54_para = para

    # 8. Replace 'Isi gambar:' editorial note with academic description
    if text.startswith("Isi gambar:"):
        if current_fig in FIG_DESC:
            set_text(para, FIG_DESC[current_fig])
            changes.append(f"  FIG  [{i+1:3d}]: Gambar 5.{current_fig} description replaced")
            if current_fig == 4:
                gambar54_has_desc = True
        current_fig = None  # consumed

    # 9. Remove internal editorial note
    if "Pada bagian implementasi sistem, gambar yang digunakan sebaiknya" in text:
        to_remove.add(id(para))
        changes.append(f"  DEL  [{i+1:3d}]: internal note paragraph")

    # 10. Remove orphaned body sentence (copy-paste artefact between figures)
    if text.strip() == (
        "Administrator memantau perubahan distribusi data "
        "dan menentukan kebutuhan retraining model."
    ):
        to_remove.add(id(para))
        changes.append(f"  DEL  [{i+1:3d}]: orphaned sentence removed")

    # 11. Find Gambar 4.2 entry inside DAFTAR GAMBAR to anchor new entries
    if "Gambar 4.2 Diagram Use Case Sistem" in text:
        gambar42_list_para = para


# ===========================================================================
# PASS 2 — remove marked paragraphs
# ===========================================================================
for para in paras:
    if id(para) in to_remove:
        remove_para(para)

# ===========================================================================
# PASS 3 — add Gambar 5.4 description if still missing
# ===========================================================================
if gambar54_para and not gambar54_has_desc:
    insert_para_after(gambar54_para, FIG_DESC[4])
    changes.append("  INS: Added Gambar 5.4 description paragraph")

# ===========================================================================
# PASS 4 — add Gambar 5.1-5.4 to DAFTAR GAMBAR
# ===========================================================================
if gambar42_list_para:
    # Insert in reverse order so final reading order is 5.1 -> 5.4
    for entry in DAFTAR_GAMBAR_NEW:
        insert_para_after(gambar42_list_para, entry, copy_style_from=gambar42_list_para)
    changes.append("  INS: Added Gambar 5.1-5.4 entries to DAFTAR GAMBAR")
else:
    changes.append("  WARN: Gambar 4.2 anchor not found — DAFTAR GAMBAR not updated")

# ===========================================================================
# Save
# ===========================================================================
doc.save(DST)
print(f"DONE. Saved: {DST}")
print(f"\nTotal change records: {len(changes)}")
for c in changes:
    print(c)
